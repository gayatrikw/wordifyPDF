# # converter/utils.py
# import io
# from docx import Document
# from PyPDF2 import PdfReader

# def convert_pdf_to_word(pdf_file):
#     pdf_content = pdf_file.read()  # Read the content of the in-memory file

#     docx = Document()
#     pdf_reader = PdfReader(io.BytesIO(pdf_content))

#     for page_number in range(len(pdf_reader.pages)):
#         page = pdf_reader.pages[page_number]
#         text = page.extract_text()
#         docx.add_paragraph(text)

#     converted_file = io.BytesIO()
#     docx.save(converted_file)
#     converted_file.seek(0)

#     return converted_file

# converter/utils.py

# ------------------------------------------------------------------------------
# Table + Images 

# import io
# import fitz  # PyMuPDF
# from docx import Document
# from docx.shared import Inches
# from PIL import Image
# import tabula

# def convert_pdf_to_word(pdf_file):
#     pdf_content = pdf_file.read()

#     docx = Document()

#     pdf_document = fitz.open("pdf", pdf_content)

#     for page_number in range(pdf_document.page_count):
#         page = pdf_document[page_number]
        
#         # Add text to Word document
#         text = page.get_text("text")
#         docx.add_paragraph(text)

#         # Add images to Word document
#         images = page.get_images(full=True)
#         for img_index, img_info in enumerate(images):
#             img_index += 1
#             image_index = img_info[0]
#             base_image = pdf_document.extract_image(image_index)
#             image_bytes = base_image["image"]

#             # # Add image to Word document
#             image_stream = io.BytesIO(image_bytes)
#             docx.add_picture(image_stream, width=Inches(4.0))

#         # Extract tables from the PDF using tabula-py
#         tables = tabula.read_pdf(pdf_file, pages=page_number + 1, multiple_tables=True)
#         for table_index, table in enumerate(tables):
#             # Add table to Word document
#             rows, cols = table.shape
#             docx.add_paragraph(f"Table {table_index + 1} - {rows} rows x {cols} columns")
#             table_data = [list(table.columns)] + table.values.tolist()
#             table_rows = len(table_data)
#             table_cols = len(table_data[0])
#             docx.add_table(rows=table_rows, cols=table_cols).style = 'Table Grid'
#             for i in range(table_rows):
#                 for j in range(table_cols):
#                     docx.tables[-1].cell(i, j).text = str(table_data[i][j])

#     # Create an in-memory buffer for the Word document
#     converted_file = io.BytesIO()
#     docx.save(converted_file)
#     converted_file.seek(0)

#     return converted_file



# ---------------------------------------------------------------------------------------




























# import io
# import fitz  # PyMuPDF
# from docx import Document
# from docx.shared import Pt
# from PIL import Image
# import tabula

# def convert_pdf_to_word(pdf_file):
#     pdf_content = pdf_file.read()

#     docx = Document()

#     pdf_document = fitz.open("pdf", pdf_content)

#     for page_number in range(pdf_document.page_count):
#         page = pdf_document[page_number]

#         # Add text to Word document
#         text = page.get_text("text")
#         paragraph = docx.add_paragraph(text)

#         # Set font size and style
#         for run in paragraph.runs:
#             font = run.font
#             font.size = Pt(12)  # Adjust the font size as needed
#             # You may need to extract more font properties based on your requirements

#         # Add images to Word document
#         images = page.get_images(full=True)
#         for img_index, img_info in enumerate(images):
#             img_index += 1
#             image_index = img_info[0]
#             base_image = pdf_document.extract_image(image_index)
#             image_bytes = base_image["image"]

#             # Add image to Word document
#             image_stream = io.BytesIO(image_bytes)
#             docx.add_picture(image_stream)

#         # Extract tables from the PDF using tabula-py
#         tables = tabula.read_pdf(pdf_file, pages=page_number + 1, multiple_tables=True)
#         for table_index, table in enumerate(tables):
#             # Add table to Word document
#             docx.add_paragraph(f"Table {table_index + 1}")
#             table_paragraph = docx.add_paragraph()
#             table_paragraph.add_run('\n'.join(['\t'.join(map(str, row)) for row in table.values]))

#     # Create an in-memory buffer for the Word document
#     converted_file = io.BytesIO()
#     docx.save(converted_file)
#     converted_file.seek(0)

#     return converted_file





# from pdf2docx import Converter

# old_pdf = "D:/Tejal/Internship-Cybrlytics/New folder (3)/Ashutosh_Arvikar.pdf"
# new_doc = "D:/Tejal/Internship-Cybrlytics/New folder (3)/Ashutosh_Arvikar_new.doc"

# obj = Converter(old_pdf)
# obj.convert(new_doc)
# obj.close()































# 99%

# import io
# from pdf2docx import Converter
# from docx import Document

# def convert_pdf_to_word(pdf_file):
#     pdf_content = pdf_file.read()

#     # Save PDF content to a temporary file
#     temp_pdf_filename = "temp.pdf"
#     with open(temp_pdf_filename, "wb") as temp_pdf_file:
#         temp_pdf_file.write(pdf_content)

#     # Convert PDF to Word using pdf2docx.Converter
#     docx_filename = "converted_document.docx"
#     cv = Converter(temp_pdf_filename)
#     cv.convert(docx_filename, start=0, end=None)
#     cv.close()

#     # Read the converted Word document
#     with open(docx_filename, "rb") as docx_file:
#         docx_content = io.BytesIO(docx_file.read())

#     return docx_content

# Example usage:
# pdf_file = open("example.pdf", "rb")
# converted_file = convert_pdf_to_word(pdf_file)
# ... (use the converted_file as needed)















import io
from pdf2docx import Converter
from docx import Document

def convert_pdf_to_word(pdf_file):
    pdf_content = pdf_file.read()

    # Save PDF content to a temporary file
    temp_pdf_filename = "temp.pdf"
    with open(temp_pdf_filename, "wb") as temp_pdf_file:
        temp_pdf_file.write(pdf_content)

    # Convert PDF to Word using pdf2docx.Converter
    docx_filename = "converted_document.docx"
    cv = Converter(temp_pdf_filename)
    cv.convert(docx_filename, start=0, end=None)
    cv.close()

    # Read the converted Word document
    with open(docx_filename, "rb") as docx_file:
        docx_content = io.BytesIO(docx_file.read())

    return docx_content

# Example usage:
# pdf_file = open("example.pdf", "rb")
# converted_file = convert_pdf_to_word(pdf_file)
# ... (use the converted_file as needed)
